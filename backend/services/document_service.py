import os
import mammoth
import logging
from docx import Document
from services.ai_service import edit_paragraph_content

logger = logging.getLogger(__name__)

def convert_doc_to_docx(input_path: str) -> str:
    """
    Converts a .doc file to .docx.
    Uses Microsoft Word via win32com if available on Windows.
    """
    if input_path.endswith(".docx"):
        return input_path
        
    output_path = input_path + "x"  # .doc -> .docx
    
    logger.info(f"[document_service] Bắt đầu chuyển đổi file .doc sang .docx: {input_path}")
    try:
        import win32com.client
        # Mở ứng dụng Word ngầm
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        # Đường dẫn bắt buộc phải là absolute
        abs_input = os.path.abspath(input_path)
        abs_output = os.path.abspath(output_path)
        
        doc = word.Documents.Open(abs_input)
        # 16 là mã định dạng cho docx
        doc.SaveAs2(abs_output, FileFormat=16)
        doc.Close()
        word.Quit()
        
        if os.path.exists(output_path):
            logger.info(f"[document_service] Chuyển đổi thành công: {output_path}")
            return output_path
        else:
            raise Exception("Chuyển đổi hoàn tất nhưng không tìm thấy file đầu ra.")
    except ImportError:
        logger.error("[document_service] Thư viện win32com không có sẵn.")
        raise Exception("Vui lòng tải lên file định dạng .docx (hoặc cài đặt pywin32 và MS Word để dùng .doc).")
    except Exception as e:
        logger.error(f"[document_service] Lỗi chuyển đổi bằng MS Word (win32com): {str(e)}")
        try:
            word.Quit()
        except:
            pass
        raise Exception("Không thể tự động chuyển đổi .doc sang .docx.")

def extract_document_elements(docx_path: str) -> dict:
    """
    Trích xuất toàn bộ các đoạn văn bản trong file Word và gán ID.
    Trả về Dictionary dạng: {"p_0": "text", "t_0_0_0_0": "text trong bảng"}
    """
    logger.info(f"[document_service] Đang bóc tách node từ {docx_path}")
    doc = Document(docx_path)
    elements = {}
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            elements[f"p_{i}"] = text
            
    for t_i, table in enumerate(doc.tables):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                for p_i, p in enumerate(cell.paragraphs):
                    text = p.text.strip()
                    if text:
                        elements[f"t_{t_i}_{r_i}_{c_i}_{p_i}"] = text
                        
    logger.info(f"[document_service] Đã bóc tách được {len(elements)} nodes.")
    return elements

def apply_node_replacements(docx_path: str, replacements_dict: dict) -> None:
    """
    Duyệt lại file Word theo đúng ID và ghi đè nội dung vào các đoạn văn tương ứng.
    """
    logger.info(f"[document_service] Bắt đầu ghi đè {len(replacements_dict)} nodes.")
    doc = Document(docx_path)
    
    for i, p in enumerate(doc.paragraphs):
        node_id = f"p_{i}"
        if node_id in replacements_dict:
            new_text = replacements_dict[node_id]
            # Cố gắng giữ lại định dạng in đậm của chữ đầu tiên nếu có
            bold = p.runs[0].bold if p.runs else None
            p.clear()
            run = p.add_run(new_text)
            if bold:
                run.bold = True
                
    for t_i, table in enumerate(doc.tables):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                for p_i, p in enumerate(cell.paragraphs):
                    node_id = f"t_{t_i}_{r_i}_{c_i}_{p_i}"
                    if node_id in replacements_dict:
                        new_text = replacements_dict[node_id]
                        bold = p.runs[0].bold if p.runs else None
                        p.clear()
                        run = p.add_run(new_text)
                        if bold:
                            run.bold = True
                            
    doc.save(docx_path)
    logger.info(f"[document_service] Ghi đè thành công.")

def process_template_nodes(file_path: str) -> tuple[str, dict]:
    """Hàm wrapper cho luồng mới"""
    docx_path = convert_doc_to_docx(file_path)
    elements = extract_document_elements(docx_path)
    return docx_path, elements

def edit_document_text_via_ai(docx_path: str, selected_text: str, paragraph_text: str, prompt: str) -> None:
    """
    Duyệt file Word, tìm đoạn văn có nội dung khớp với paragraph_text,
    gọi AI để sửa và ghi đè lại file.
    """
    logger.info(f"[document_service] Tìm và sửa đoạn văn. File: {docx_path}")
    doc = Document(docx_path)
    
    sel_clean = "".join(selected_text.split())
    para_clean = "".join(paragraph_text.split())
    
    if not sel_clean:
        raise Exception("Nội dung bôi đen trống.")
        
    candidates = []
    
    def extract_candidates(paragraphs):
        for p in paragraphs:
            if p.text.strip():
                p_clean = "".join(p.text.split())
                if para_clean in p_clean or p_clean in para_clean:
                    candidates.append(p)
                    
    extract_candidates(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_candidates(cell.paragraphs)
                
    target_p = None
    if len(candidates) >= 1:
        for p in candidates:
            p_clean = "".join(p.text.split())
            if p_clean == para_clean:
                target_p = p
                break
        if not target_p:
            target_p = candidates[0]
            
    if target_p:
        logger.info("[document_service] Đã tìm thấy đoạn văn cần sửa.")
        new_text = edit_paragraph_content(target_p.text, selected_text, prompt)
        bold = target_p.runs[0].bold if target_p.runs else None
        target_p.clear()
        
        run = target_p.add_run()
        if bold:
            run.bold = True
            
        lines = new_text.split('\n')
        for i, line in enumerate(lines):
            tab_parts = line.split('\t')
            for j, t_part in enumerate(tab_parts):
                run.add_text(t_part)
                if j < len(tab_parts) - 1:
                    run.add_tab()
            if i < len(lines) - 1:
                run.add_break()
                
        doc.save(docx_path)
        logger.info(f"[document_service] Đã lưu file sau khi sửa: {docx_path}")
    else:
        logger.warning("[document_service] Không tìm thấy đoạn văn khớp với nội dung đã bôi đen.")
        raise Exception("Không tìm thấy đoạn văn tương ứng trong file Word.")

def edit_document_text_manual(docx_path: str, paragraph_text: str, new_paragraph_text: str) -> None:
    """
    Duyệt file Word, tìm đoạn văn có nội dung khớp với paragraph_text,
    và ghi đè bằng new_paragraph_text (sửa thủ công).
    """
    logger.info(f"[document_service] Tìm và sửa đoạn văn (thủ công). File: {docx_path}")
    doc = Document(docx_path)
    
    para_clean = "".join(paragraph_text.split())
    
    if not para_clean:
        raise Exception("Nội dung đoạn văn gốc trống.")
        
    candidates = []
    
    def extract_candidates(paragraphs):
        for p in paragraphs:
            if p.text.strip():
                p_clean = "".join(p.text.split())
                if para_clean in p_clean or p_clean in para_clean:
                    candidates.append(p)
                    
    extract_candidates(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_candidates(cell.paragraphs)
                
    target_p = None
    if len(candidates) >= 1:
        # Lấy cái khớp nhất hoặc cái đầu tiên
        for p in candidates:
            p_clean = "".join(p.text.split())
            if p_clean == para_clean:
                target_p = p
                break
        if not target_p:
            target_p = candidates[0]
            
    if target_p:
        logger.info("[document_service] Đã tìm thấy đoạn văn cần sửa thủ công.")
        bold = target_p.runs[0].bold if target_p.runs else None
        target_p.clear()
        
        run = target_p.add_run()
        if bold:
            run.bold = True
            
        # Parse newlines and tabs
        lines = new_paragraph_text.split('\n')
        for i, line in enumerate(lines):
            tab_parts = line.split('\t')
            for j, t_part in enumerate(tab_parts):
                run.add_text(t_part)
                if j < len(tab_parts) - 1:
                    run.add_tab()
            if i < len(lines) - 1:
                run.add_break()
                
        doc.save(docx_path)
        logger.info(f"[document_service] Đã lưu file sau khi sửa thủ công: {docx_path}")
    else:
        logger.warning("[document_service] Không tìm thấy đoạn văn gốc để sửa thủ công.")
        raise Exception("Không tìm thấy đoạn văn tương ứng trong file Word để sửa.")
