from docx import Document
import json
import os

def test_extract():
    docx_path = os.path.abspath(r"temp\Phu_luc_4_Mau_Ke_hoach_bai_day.docx")
    doc = Document(docx_path)
    elements = {}
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            elements[f"p_{i}"] = text
            
    for t_i, table in enumerate(doc.tables):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                for p_i, p in enumerate(cell.paragraphs):
                    text = p.text.strip()
                    if text:
                        elements[f"t_{t_i}_{r_i}_{c_i}_{p_i}"] = text
                        
    print(f"Extracted {len(elements)} elements.")
    print("Sample:")
    for k in list(elements.keys())[:10]:
        print(f"  {k}: {elements[k]}")

if __name__ == "__main__":
    test_extract()
