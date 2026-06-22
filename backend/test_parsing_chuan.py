from docx import Document
import os
import json

def test_extract():
    docx_path = os.path.abspath(r"templates\mau_giao_an_chuan.docx")
    doc = Document(docx_path)
    elements = {}
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            elements[f"p_{i}"] = text
            
    for t_i, table in enumerate(doc.tables):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                for p_i, p in enumerate(cell.paragraphs):
                    text = p.text.strip()
                    if text:
                        elements[f"t_{t_i}_{r_i}_{c_i}_{p_i}"] = text
                        
    with open("elements_dump.json", "w", encoding="utf-8") as f:
        json.dump(elements, f, ensure_ascii=False, indent=2)
        
    print("Dumped elements to elements_dump.json")

if __name__ == "__main__":
    test_extract()
