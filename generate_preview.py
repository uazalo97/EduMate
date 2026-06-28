import os
from docx import Document
import re

def replace_text_in_paragraph(paragraph, replacements):
    text = paragraph.text
    for key, val in replacements.items():
        if key in text:
            text = text.replace(key, val)
    
    if text != paragraph.text:
        # Clear all runs and add a single run with the new text, preserving the alignment
        paragraph.clear()
        paragraph.add_run(text)

def generate_preview():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(base_dir, "backend", "templates", "mau_giao_an_chuan.docx")
    out_path = os.path.join(base_dir, "backend", "templates", "mau_giao_an_chuan_preview.docx")
    
    doc = Document(template_path)
    
    replacements = {
        "{{ ten_truong }}": "Trường THPT Nguyễn Trãi",
        "{{ ten_to }}": "Tổ Khoa học Tự nhiên",
        "{{ ten_giao_vien }}": "Nguyễn Văn A",
        "{{ ten_bai_day }}": "BÀI 1: KHÁI QUÁT VỀ... (MẪU XEM TRƯỚC)",
        "{{ mon_hoc }}": "Vật lý",
        "{{ lop }}": "10",
        "{{ so_tiet }}": "2",
        "{{ kien_thuc }}": "- Nêu được các khái niệm cơ bản...\n- Trình bày được định luật...",
        "{{ nang_luc }}": "- Năng lực giải quyết vấn đề...\n- Năng lực tính toán...",
        "{{ pham_chat }}": "- Chăm chỉ học tập, rèn luyện...\n- Có trách nhiệm trong hoạt động nhóm...",
        "{{ thiet_bi_day_hoc_va_hoc_lieu }}": "- Sách giáo khoa, máy chiếu, bài giảng PowerPoint.\n- Phiếu học tập, dụng cụ thí nghiệm.",
        "{{ tieu_de_hoat_dong_1 }}": "1. Hoạt động 1: Khởi động (Mở đầu)",
        "{{ muc_tieu_1 }}": "Tạo tâm thế hứng thú, khơi gợi kiến thức nền của học sinh.",
        "{{ noi_dung_1 }}": "Giáo viên chiếu video về hiện tượng... và đặt câu hỏi mở.",
        "{{ san_pham_1 }}": "Câu trả lời dự kiến của học sinh dựa trên video.",
        "{{ to_chuc_thuc_hien_1 }}": "Bước 1: GV giao nhiệm vụ...\nBước 2: HS thảo luận nhóm...\nBước 3: Đại diện nhóm trình bày...",
        "{{ tieu_de_hoat_dong_2 }}": "2. Hoạt động 2: Hình thành kiến thức mới",
        "{{ muc_tieu_2 }}": "Học sinh nắm được khái niệm cốt lõi và công thức...",
        "{{ noi_dung_2 }}": "Giáo viên giảng giải kết hợp sơ đồ tư duy...",
        "{{ san_pham_2 }}": "Bản ghi chép của học sinh, sơ đồ tư duy hoàn thiện.",
        "{{ to_chuc_thuc_hien_2 }}": "Hoạt động nhóm đôi nghiên cứu tài liệu và hoàn thiện phiếu học tập.",
        "{{ tieu_de_hoat_dong_3 }}": "3. Hoạt động 3: Luyện tập",
        "{{ muc_tieu_3 }}": "Củng cố kiến thức vừa học, áp dụng giải bài tập cơ bản.",
        "{{ noi_dung_3 }}": "Hệ thống bài tập trắc nghiệm và tự luận nhanh.",
        "{{ san_pham_3 }}": "Đáp án bài tập trong vở học sinh.",
        "{{ to_chuc_thuc_hien_3 }}": "Thi đấu giữa các tổ, chấm điểm chéo.",
        "{{ tieu_de_hoat_dong_4 }}": "4. Hoạt động 4: Vận dụng",
        "{{ muc_tieu_4 }}": "Vận dụng kiến thức vào thực tiễn cuộc sống.",
        "{{ noi_dung_4 }}": "Bài tập dự án nhỏ về nhà.",
        "{{ san_pham_4 }}": "Báo cáo dự án nhóm vào tuần sau.",
        "{{ to_chuc_thuc_hien_4 }}": "Giáo viên hướng dẫn chi tiết yêu cầu, tiêu chí đánh giá.",
    }
    
    # Process paragraphs
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)
        
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements)
                    
    doc.save(out_path)
    print(f"Generated {out_path}")

if __name__ == '__main__':
    generate_preview()
